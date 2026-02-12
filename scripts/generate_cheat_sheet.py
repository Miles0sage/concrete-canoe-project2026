#!/usr/bin/env python3
"""
Generate a one-page engineering cheat sheet for Design A.
All 7 formulas + ASCE checks on a single printable page.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
NAVY = RGBColor(0x2E, 0x40, 0x57)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x15, 0x5C, 0x24)


def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'})
    tcPr.append(shd)


def add_formula_block(doc, num, title, question, formula_lines, design_a_calc, answer, asce_check=None):
    """Add a compact formula block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(f"{num}. {title}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY

    run = p.add_run(f"  —  {question}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.italic = True

    # Formula
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    for line in formula_lines:
        run = p.add_run(line + "\n")
        run.font.size = Pt(7.5)
        run.font.name = "Courier New"
        run.font.color.rgb = DARK

    # Design A calculation
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("Design A: ")
    run.bold = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = DARK
    run = p.add_run(design_a_calc)
    run.font.size = Pt(7.5)
    run.font.color.rgb = DARK

    # Answer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"= {answer}")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = GREEN

    if asce_check:
        run = p.add_run(f"    {asce_check}")
        run.font.size = Pt(7.5)
        run.font.color.rgb = GREEN
        run.bold = True


def main():
    doc = Document()

    # Tight margins for one-page fit
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CONCRETE CANOE ENGINEERING CHEAT SHEET")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Design A: 192\" x 32\" x 17\" x 0.5\" wall  |  174.26 lbs  |  60 PCF  |  f'r = 1500 psi  |  Crew: 700 lbs")
    run.font.size = Pt(7.5)
    run.font.color.rgb = GRAY

    # ── Formula 1: Hull Weight ──
    add_formula_block(doc, 1, "HULL WEIGHT", "How heavy is it?",
        [
            "girth = B + 2D    shell_area = girth x L x Cp    vol = area x t",
            "weight = vol x density x 1.10   (Cp=0.55 taper, 1.10 overhead)",
        ],
        "girth=5.50 ft, area=48.4 ft², vol=2.02 ft³, est=133 lbs. Actual = shell(163.11) + reinf(8.16) + finish(3.00)",
        "174.26 lbs  (target <= 237 lbs)"
    )

    # ── Formula 2: Draft ──
    add_formula_block(doc, 2, "DRAFT (Archimedes)", "How deep does it sit?",
        [
            "V_disp = W_total / 62.4        (Archimedes: weight = displaced water)",
            "A_wp = L x B x Cwp             (Cwp=0.70 canoe waterplane coefficient)",
            "Draft = V_disp / A_wp",
        ],
        "V=874.26/62.4=14.01 ft³, A_wp=16.0x2.667x0.70=29.87 ft², T=14.01/29.87=0.469 ft",
        "Draft = 5.63 in"
    )

    # ── Formula 3: Freeboard ──
    add_formula_block(doc, 3, "FREEBOARD", "How much hull above water?",
        ["Freeboard = Depth - Draft"],
        "F = 17.0 - 5.63",
        "Freeboard = 11.37 in",
        "PASS (ASCE req >= 6.0\")"
    )

    # ── Formula 4: Stability GM ──
    add_formula_block(doc, 4, "STABILITY (GM)", "Will it tip over?",
        [
            "GM = KB + BM - KG",
            "KB = T/2    BM = I_wp/V_disp    I_wp = Cwp x L x B^3 / 12",
            "KG = (W_hull x h_hull + W_crew x h_crew) / W_total",
        ],
        "KB=2.81\", BM=I_wp/V_disp=15.16\", KG=(174.26x0.539+700x0.833)/874.26=0.774ft=9.29\"",
        "GM = 2.81 + 15.16 - 9.29 = 8.68 in",
        "PASS (ASCE req >= 6.0\")"
    )

    # ── Formula 5: Bending Moment ──
    add_formula_block(doc, 5, "BENDING MOMENT", "How much does it want to snap?",
        [
            "M_hull = (W/L) x L^2 / 8       (uniform distributed load, SS beam)",
            "M_crew = P x L / 4              (point load at midspan)",
            "M_total = M_hull + M_crew",
        ],
        "w=174.26/16=10.89 lb/ft, M_hull=10.89x16²/8=348.5, M_crew=700x16/4=2800.0",
        "M_total = 3148.5 lb-ft"
    )

    # ── Formula 6: Section Modulus ──
    add_formula_block(doc, 6, "SECTION MODULUS", "How strong is the cross-section?",
        [
            "U-shape: bottom plate (BxT) + 2 side walls (Tx(D-T))",
            "y_NA = sum(A_i x y_i) / A_total    (neutral axis location)",
            "I = sum(I_self + A x d^2)           (parallel axis theorem)",
            "S = (I / c_max) x 0.75              (ACI thin-shell reduction)",
        ],
        "A_bot=16.0, A_wall=8.25x2, y_NA=5.10\", I=412.5 in^4, c_max=11.90\"",
        "S = 58.00 in^3  (NOT bh^2/6=1541 -- that's solid rect, 27x too high!)"
    )

    # ── Formula 7: Safety Factor ──
    add_formula_block(doc, 7, "SAFETY FACTOR", "Will it break?",
        [
            "sigma = M x 12 / S              (bending stress, convert lb-ft to lb-in)",
            "SF = f'r / sigma                 (strength / stress)",
        ],
        "sigma = 3148.5 x 12 / 58.00 = 651.45 psi, SF = 1500 / 651.45",
        "SF = 2.30",
        "PASS (ASCE req >= 2.0)"
    )

    # ── ASCE Compliance Summary Table ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ASCE 2026 COMPLIANCE SUMMARY")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ["Requirement", "Threshold", "Design A", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E4057')

    data = [
        ("Freeboard", ">= 6.0 in", "11.37 in", "PASS"),
        ("Metacentric Height (GM)", ">= 6.0 in", "8.68 in", "PASS"),
        ("Safety Factor", ">= 2.0", "2.30", "PASS"),
        ("Portland Cement", "<= 40% of CM", "30.8%", "PASS"),
        ("Water/Cementitious", "<= 0.50", "0.36", "PASS"),
        ("Reinforcement", "<= 50% of wall", "0.34%", "PASS"),
    ]

    for r, (req, thresh, val, status) in enumerate(data):
        row = table.rows[r + 1]
        row.cells[0].text = req
        row.cells[1].text = thresh
        row.cells[2].text = val
        row.cells[3].text = status
        for c in range(4):
            for p in row.cells[c].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7.5)
        if status == "PASS":
            set_cell_shading(row.cells[3], 'D4EDDA')

    # ── Footer ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(
        "References: Archimedes (287 BC) | Bouguer (1746) | ACI 318-19 | "
        "Beer & Johnston Mechanics of Materials | Hibbeler Structural Analysis | "
        "SNAME Principles of Naval Architecture"
    )
    run.font.size = Pt(6.5)
    run.font.color.rgb = GRAY
    run.italic = True

    output = PROJECT_ROOT / "reports" / "Engineering_Cheat_Sheet.docx"
    doc.save(str(output))
    print(f"Cheat sheet saved: {output}")
    return output


if __name__ == "__main__":
    main()
